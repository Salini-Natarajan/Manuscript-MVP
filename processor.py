import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
import time

# Configure Gemini
genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
model = genai.GenerativeModel('gemini-1.5-flash')

def classify_text(text):
    """Sends text to Gemini to classify it for specific formatting rules."""
    prompt = f"""
    Analyze this text from a raw manuscript and classify it into exactly ONE of these categories:
    - Title
    - Heading 1
    - Heading 2
    - Paragraph
    - Figure Caption (Text describing an image, usually starting with 'Fig' or 'Figure')
    - Table Caption (Text describing a table, usually starting with 'Table')

    Text to analyze: "{text}"
    
    Respond with ONLY the exact category name and nothing else.
    """
    try:
        response = model.generate_content(prompt)
        return response.text.strip().replace(".", "")
    except Exception:
        return "Paragraph"

def apply_custom_formatting(paragraph, classification):
    """Applies your exact font, size, alignment, and spacing rules directly to the text."""
    
    # Requirement 1 & 2: Line spacing 1.5 AND Add spaces before paragraph
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(0)
    
    # Default parameters
    align = WD_ALIGN_PARAGRAPH.LEFT
    font_size = Pt(12)
    bold = False
    italic = False
    color = RGBColor(0, 0, 0) # Black

    # Apply specific rules based on AI classification
    if classification == "Title":
        align = WD_ALIGN_PARAGRAPH.CENTER
        font_size = Pt(24) 
        bold = True
    
    elif classification == "Heading 1":
        font_size = Pt(14) 
        bold = True
    
    elif classification == "Heading 2":
        font_size = Pt(13) 
        bold = True
        color = RGBColor(0, 0, 255) # Blue
        
    elif classification == "Paragraph":
        align = WD_ALIGN_PARAGRAPH.JUSTIFY 
        
    elif classification in ["Figure Caption", "Table Caption"]:
        # Requirement 4 & 5: Center align captions and italicize them
        align = WD_ALIGN_PARAGRAPH.CENTER 
        font_size = Pt(11)
        italic = True 

    # Apply alignment to paragraph
    paragraph.alignment = align

    # Apply font rules to every word (Run) in the paragraph
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = font_size
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color

def get_word_count(uploaded_file):
    """Scans the document and returns the total word count."""
    doc = Document(uploaded_file)
    full_text = []
    
    for para in doc.paragraphs:
        full_text.append(para.text)
        
    # Join all text and count the words
    total_words = len(" ".join(full_text).split())
    return total_words

def process_document(uploaded_file, style):
    doc = Document(uploaded_file)
    
    progress_text = "AI is analyzing and enforcing custom formatting..."
    progress_bar = st.progress(0, text=progress_text)
    
    # Requirement 3: Adjust scale (Margins) to standardize the layout
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Requirement 5 (Part 1): Center align all Tables
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

    total_paras = len(doc.paragraphs)
    
    for i, para in enumerate(doc.paragraphs):
        # Requirement 4 (Part 1): Center align Figures
        if 'graphicData' in para._p.xml:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        text = para.text.strip()
        if text != "":
            # Ask Gemini what this text is
            classification = classify_text(text)
            
            # Force the exact formatting rules onto the text
            apply_custom_formatting(para, classification)
            
            # API Rate Limit Protection
            time.sleep(2) 
            
        progress_bar.progress((i + 1) / total_paras, text=f"Processed paragraph {i+1} of {total_paras}")
        
    progress_bar.empty()
    
    # Save output securely
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output